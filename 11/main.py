import openpyxl
import json
import requests
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
data1 = {'Data': [1111]}
data2 = {'Data': [2222]}
data3 = {'Data': [3333]}
wb1 = openpyxl.Workbook()
ws1 = wb1.active
ws1.append(data1['Data'])
wb1.save('file1.xlsx')

wb2 = openpyxl.Workbook()
ws2 = wb2.active
ws2.append(data2['Data'])
wb2.save('file2.xlsx')

wb3 = openpyxl.Workbook()
ws3 = wb3.active
ws3.append(data3['Data'])
wb3.save('file3.xlsx')
files = ['file1.xlsx', 'file2.xlsx', 'file3.xlsx']
merged_data = []
for file in files:
    wb = openpyxl.load_workbook(file)
    ws = wb.active
    merged_data.extend(ws.values)

sorted_data = sorted(merged_data, key=lambda x: x[0], reverse=True)
wb_result = openpyxl.Workbook()
ws_result = wb_result.active
for row in sorted_data:
    ws_result.append(row)
wb_result.save('result.xlsx')



url = 'https://jsonplaceholder.typicode.com/todos/'
response = requests.get(url)
todos_data = response.json()
with open('todos.json', 'w') as json_file:
    json.dump(todos_data, json_file)
with open('todos.json', 'r') as json_file:
    todos_list = json.load(json_file)
for i, todo in enumerate(todos_list):
    with open(f'todo_{i + 1}.json', 'w') as json_file:
        json.dump(todo, json_file)
        


doc = Document()
doc.add_paragraph("Hello Python")
doc.save('original.docx')

doc = Document('original.docx')
bold_text = ''
for paragraph in doc.paragraphs:
    if paragraph.runs and paragraph.runs[0].bold:
        bold_text += paragraph.text
print(bold_text)
new_doc = Document()
new_paragraph = new_doc.add_paragraph("New paragraph text.")
new_run = new_paragraph.runs[0]
new_run.bold = True
new_run.font.size = Pt(14)
new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

new_doc.save('new_file.docx')
